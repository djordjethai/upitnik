import matplotlib.pyplot as plt
import numpy as np
import os
import streamlit as st

from docx import Document
from docx.shared import Pt
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from openai import OpenAI
from pitanja import odgovori
from smtplib import SMTP

from myfunc.retrievers import HybridQueryProcessor

#from myfunc.mojafunkcija import send_email

client=OpenAI()
avatar_ai="bot.png" 


# docx koji sadrzi formatirani tekst (i grafikon)
def create_docx(content, image_path=None, filename="formatted_document.docx"):
    doc = Document()
    # Define custom styles for document if not already present
    styles = doc.styles
    if 'Heading 1' not in styles:
        heading1 = styles.add_style('Heading 1', 1)
        heading1.font.size = Pt(24)

    if 'Heading 2' not in styles:
        heading2 = styles.add_style('Heading 2', 1)
        heading2.font.size = Pt(20)

    # Parse Markdown and apply styles
    lines = content.split('\n')
    for line in lines:
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 1')
        elif line.startswith('#### '):
            p = doc.add_paragraph(line[5:], style='Heading 2')
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif '**' in line:
            # Handle bold within line
            p = doc.add_paragraph()
            parts = line.split('**')
            bold = False
            for part in parts:
                run = p.add_run(part)
                if bold:
                    run.bold = True
                bold = not bold
        else:
            p = doc.add_paragraph(line)

    # Add image if provided
    if image_path:
        doc.add_picture(image_path, width=Pt(300))  # Width is just an example

    # Save the document
    doc.save(filename)
    return filename


# salje mejl
def posalji_mail(email, gap_analiza, image_path):
    st.info(f"Saljem email na adresu {email}")
    doc_path = create_docx(gap_analiza, image_path)  # Now also passing the image path

    # Assuming your send_email function can handle attachments and is defined as shown previously
    send_email(
        subject="Gap Analiza",
        message="Please find attached the gap analysis document, which includes the radar chart.",
        from_addr="azure.test@positive.rs",
        to_addr=email,
        smtp_server="smtp.office365.com",
        smtp_port=587,
        username="azure.test@positive.rs",
        password=os.getenv("PRAVNIK_PASS"),
        attachments=[doc_path]  # Only the document needs to be attached now, as it includes the image
    )
    st.info(f"Poslat je email na adresu {email}")


# Slanje mejla sa attachmentom
def send_email(subject, message, from_addr, to_addr, smtp_server, smtp_port, username, password, attachments=None):
    """
    Sends an email using SMTP protocol.

    Args:
    subject (str): Subject line of the email.
    message (str): Body of the email.
    from_addr (str): Sender's email address.
    to_addr (str): Recipient's email address.
    smtp_server (str): Address of the SMTP server to connect to.
    smtp_port (int): Port number for the SMTP server.
    username (str): Username for the SMTP server authentication.
    password (str): Password for the SMTP server authentication.
    attachments (list of str): List of file paths to attach to the email.
    
    This function creates an email message using the specified subject and
    message, sets up a connection to the specified SMTP server, logs in with
    provided credentials, and sends the email. The connection is securely 
    established using TLS (Transport Layer Security).
    """
    
    msg = MIMEMultipart()
    msg['From'] = from_addr
    msg['To'] = to_addr
    msg['Subject'] = subject

    msg.attach(MIMEText(message, 'plain'))  # Attach the body text
    
    if attachments:  # Attach any files specified
        for attachment in attachments:
            part = MIMEBase('application', "octet-stream")
            with open(attachment, 'rb') as file:
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment)}')
            msg.attach(part)

    # Connect to the SMTP server and send the email
    server = SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(username, password)
    server.send_message(msg)  # Recommended to use send_message() as it handles message encoding issues.
    server.quit()


# agent llm odgovara na razlicite upite - treba u myfunc
def positive_agent(messages):
    with st.chat_message("assistant", avatar=avatar_ai):
        message_placeholder = st.empty()
        full_response = ""
        for response in client.chat.completions.create(
            model="gpt-4-turbo-preview",
            messages=messages,
            stream=True,
        ):
            full_response += (response.choices[0].delta.content or "")
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)
        
    return full_response


# privremeni grafikon
def create_radar_chart(data, labels, num_vars):
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    data += data[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, data, color='red', alpha=0.25)
    ax.plot(angles, data, color='red', linewidth=2)  # Draw the outline
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    return fig


# cuva grafikon
def show_graph():
    st.title('Polar Chart Example')
    labels = ['Poslovna zrelostt', 'Digitalna zrelost', 'Upotreba AI', 'Sajber bezbednost', 'IT infrastruktura']
    data = [4, 3, 4, 2, 5]
    num_vars = len(data)
    # Plotting
    radar_fig = create_radar_chart(data, labels, num_vars)
    radar_fig.savefig('radar_chart.png', bbox_inches='tight')
    # Display in Streamlit
    st.pyplot(radar_fig)
    return 'radar_chart.png'

 
# RAG pretrazuje index za preporuke    
def recommended(full_response):
    processor = HybridQueryProcessor(namespace="positive", top_k=3)
    return processor.process_query_results(full_response)


# glavni program
def main():
    with st.sidebar:
        st.caption("Ver. 14.04.24" )
        st.subheader("Demo GAP sa grafikonon i slanjem maila ")
        opcija = st.selectbox("Odaberite upitnik", ("",
                                                    "Opsti", 
                                                    "Poslovna zrelost", 
                                                    "Digitalna zrelost", 
                                                    "Sajber bezbednost", 
                                                    "IT infrastruktura", 
                                                    "Upotreba AI" ))
    if opcija !="":  # Check if the result is not None
        result, email = odgovori(opcija)
        if result:
            # prva faza citanje odgovora i komentar
            gap_message=[
                {"role": "system", "content": """[Use only Serbian language] You are an expert in business data analysis. \
                 Analyze the document. Think critically and do business analysis of the company. The accent is on GAP analysis, focusing on 
                 generating sections Current state, Desired state. """},

                {"role": "user", "content": f"Write your GAP analysis report based on this input: {result}"}
            ]
            full_response = positive_agent(gap_message)
            predlozi = recommended(full_response)
            # druga faza preporuke na osnovu portfolia
            recommend_message=[
                        {"role": "system", "content": """[Use only Serbian Language] \
                         You are an experienced digital transformation consultant. \
                         You are working for company Positive doo, the leader in Digital Transformation services in Serbia.
                         When making propositions, convince the client in a non-invasive way that hiring your company is the right choice."""},

                        {"role": "user", "content": f"""Based on previous GAP analysis: {full_response}, \
                         make suggestions for business improvement of the descibed business process. \
                         Write in potential, not in the future tense.
                         Suggest solutions in the form of the proposal (offer) based on the text from portfolio of your company Positive doo: {predlozi}!!!"""}
            ]
            recommendation_response = positive_agent(recommend_message)    
            # treca faza kreiranje dokumenta
            grafikon = show_graph()
            gap_analiza = full_response + "\n\n" + recommendation_response + "\n\n"
            # cetvrta faza slanje maila
            posalji_mail(email, gap_analiza, grafikon)
            os.remove("formatted_document.docx")
                
if __name__ == "__main__":
    main()