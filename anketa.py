import os
import streamlit as st
import json
from docx import Document
from docx.shared import Pt
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pitanja import odgovori
from smtplib import SMTP
#
# program sakuplja odgovore ankete i salje ih kao word fajl n aodabrani email
def format_json_to_text(data):
    output = []
    for key, value in data.items():
        if isinstance(value, list):  # If the value is a list, format it into a comma-separated string
            value = ", ".join(value)
        # Append the formatted string to the output list
        output.append(f"{key}: '\n' {value}")
    # Join all items in the output list with newlines to create the final string
    return "\n\n".join(output)

# docx koji sadrzi formatirani tekst (i grafikon)
def create_docx(content, image_path=None, filename="Anketa.docx"):
    doc = Document()
    # Define custom styles for document if not already present
    styles = doc.styles
    if 'Heading 1' not in styles:
        heading1 = styles.add_style('Heading 1', 1)
        heading1.font.size = Pt(24)

    if 'Heading 2' not in styles:
        heading2 = styles.add_style('Heading 2', 1)
        heading2.font.size = Pt(20)

    # Parse Markdown and apply styles
    lines = content.split('\n')
    for line in lines:
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 1')
        elif line.startswith('#### '):
            p = doc.add_paragraph(line[5:], style='Heading 2')
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif '**' in line:
            # Handle bold within line
            p = doc.add_paragraph()
            parts = line.split('**')
            bold = False
            for part in parts:
                run = p.add_run(part)
                if bold:
                    run.bold = True
                bold = not bold
        else:
            p = doc.add_paragraph(line)

    # Add image if provided
    if image_path:
        doc.add_picture(image_path, width=Pt(300))  # Width is just an example

    # Save the document
    doc.save(filename)
    return filename


# salje mejl
def posalji_mail(email, gap_analiza, image_path, filename="Anketa.docx"):
    st.info(f"Šaljem email na adresu {email}")
    doc_path = create_docx(gap_analiza, image_path, filename)  # Now also passing the image path

    # Assuming your send_email function can handle attachments and is defined as shown previously
    send_email(
        subject="AI upitnik",
        message="Odgovor se nalaze u prilogu.",
        from_addr="azure.test@positive.rs",
        to_addr="djordje.medakovic@positive.rs",
        smtp_server="smtp.office365.com",
        smtp_port=587,
        username="azure.test@positive.rs",
        password=os.getenv("PRAVNIK_PASS"),
        attachments=[doc_path]  # Only the document needs to be attached now, as it includes the image
    )
    send_email(
        subject="AI upitnik",
        message="Please find attached AI questionnaire.",
        from_addr="azure.test@positive.rs",
        to_addr=email,
        smtp_server="smtp.office365.com",
        smtp_port=587,
        username="azure.test@positive.rs",
        password=os.getenv("PRAVNIK_PASS"),
        attachments=[doc_path]  # Only the document needs to be attached now, as it includes the image
    )
    st.info(f"Poslat je email na adresu {email}")


# Slanje mejla sa attachmentom
def send_email(subject, message, from_addr, to_addr, smtp_server, smtp_port, username, password, attachments=None):
    """
    Sends an email using SMTP protocol.

    Args:
    subject (str): Subject line of the email.
    message (str): Body of the email.
    from_addr (str): Sender's email address.
    to_addr (str): Recipient's email address.
    smtp_server (str): Address of the SMTP server to connect to.
    smtp_port (int): Port number for the SMTP server.
    username (str): Username for the SMTP server authentication.
    password (str): Password for the SMTP server authentication.
    attachments (list of str): List of file paths to attach to the email.
    
    This function creates an email message using the specified subject and
    message, sets up a connection to the specified SMTP server, logs in with
    provided credentials, and sends the email. The connection is securely 
    established using TLS (Transport Layer Security).
    """
    
    msg = MIMEMultipart()
    msg['From'] = from_addr
    msg['To'] = to_addr
    msg['Subject'] = subject

    msg.attach(MIMEText(message, 'plain'))  # Attach the body text
    
    if attachments:  # Attach any files specified
        for attachment in attachments:
            part = MIMEBase('application', "octet-stream")
            with open(attachment, 'rb' ) as file:
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment)}')
            msg.attach(part)

    # Connect to the SMTP server and send the email
    server = SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(username, password)
    server.send_message(msg)  # Recommended to use send_message() as it handles message encoding issues.
    server.quit()

# glavni program
def main():
    filename = "Anketa_popunjena.docx"
    with st.sidebar:
        st.caption("Ver. 17.04.24" )
        st.subheader("Demo Anketa i slanje maila ")
        opcija = "Anketa" 
    if opcija !="":  # Check if the result is not None
        result, email = odgovori(opcija)
        if result:
            gap_analiza = format_json_to_text(result)
            # slanje maila
            posalji_mail(email, gap_analiza, None, filename)
            os.remove(filename)
                
if __name__ == "__main__":
    main()