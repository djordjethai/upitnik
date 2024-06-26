﻿import os
import streamlit as st
import subprocess

from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from openai import OpenAI
from pitanja import odgovori
from smtplib import SMTP

from myfunc.retrievers import HybridQueryProcessor
from myfunc.varvars_dicts import work_prompts, work_vars

mprompts = work_prompts()

client = OpenAI()
avatar_ai = "bot.png"
anketa = ""

def change_extension(filename, new_extension):
    base = os.path.splitext(filename)[0]
    return base + new_extension

def format_json_to_text(data):
    output = []
    for key, value in data.items():
        if isinstance(value, list):  
            value = ", ".join(value)
        output.append(f"{key} '\n\n' {value}")
    return "\n\n".join(output)

def add_markdown_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    bold = False
    parts = text.split('**')
    for part in parts:
        run = p.add_run(part)
        if bold:
            run.bold = True
        bold = not bold

    if style and (style.startswith('Heading 1') or style.startswith('Heading 2')):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def convert_docx_to_pdf(docx_file_path, pdf_file_path):
    # Use LibreOffice to convert the DOCX file to PDF
    result = subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf:writer_pdf_Export', docx_file_path, '--outdir', os.path.dirname(pdf_file_path)], check=True)
    if result.returncode != 0:
        raise subprocess.CalledProcessError(result.returncode, result.args)

def sacuvaj_dokument_upitnik(content, file_name, template_path="template.docx", anketa=anketa):
    if template_path:
        doc = Document(template_path)
    else:
        doc = Document()

    datum = datetime.today().date()
    formatted_date = datum.strftime("%d.%m.%Y.")

    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            add_markdown_paragraph(doc, line[2:], style='Heading 1')
        elif line.startswith('## '):
            add_markdown_paragraph(doc, line[3:], style='Heading 2')
        elif line.startswith('### '):
            add_markdown_paragraph(doc, line[4:], style='Heading 3')
        elif line.startswith('#### '):
            add_markdown_paragraph(doc, line[5:], style='Heading 4')
        elif line.startswith('- ') or line.startswith(' - ') or line.startswith('  - ') or line.startswith('   - ') or line.startswith(' -  ') or line.startswith('  -  '):
            add_markdown_paragraph(doc, line[2:], style='List Paragraph')
        else:
            add_markdown_paragraph(doc, line)
    doc.add_paragraph(" ")
    doc.add_paragraph(f"Datum {formatted_date}")
    doc.add_page_break()
    doc.add_paragraph(f"Anketa\n\n", style='Heading 2')
    lines = anketa.split('\n')
    for line in lines:
        try:
            doc.add_paragraph(line)
        except:
            pass
    doc.save(file_name)

    pdf_file_name = change_extension(file_name, ".pdf")
    convert_docx_to_pdf(file_name, pdf_file_name)
    return pdf_file_name

def posalji_mail(email, file_name, poruka):
    st.info(f"Sending email to {email}")
    cwd = os.getcwd()
    file_path = os.path.join(cwd, file_name)

    send_email(
        subject="Izveštaj - Gap Analiza",
        message=poruka,
        from_addr="Aiupitnik@positive.rs",
        to_addrs=[email, "prodaja@positive.rs"],
        smtp_server="smtp.office365.com",
        smtp_port=587,
        username="Aiupitnik@positive.rs",
        password=os.getenv("PRAVNIK_PASS"),
        attachments=[file_path]
    )
    st.info(f"Email sent to {email}")
    os.remove(file_path)

def send_email(subject, message, from_addr, to_addrs, smtp_server, smtp_port, username, password, attachments=None):
    for to_addr in to_addrs:
        msg = MIMEMultipart()
        msg['From'] = from_addr
        msg['To'] = to_addr
        msg['Subject'] = subject
        msg.attach(MIMEText(message, 'plain'))

        if attachments:
            for attachment in attachments:
                with open(attachment, 'rb') as file:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(file.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment))
                msg.attach(part)

        server = SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(username, password)
        server.send_message(msg)
        server.quit()

def positive_agent(messages):
    with st.chat_message("assistant", avatar=avatar_ai):
        message_placeholder = st.empty()
        full_response = ""
        for response in client.chat.completions.create(
            model=work_vars["names"]["openai_model"],
            messages=messages,
            stream=True,
        ):
            full_response += (response.choices[0].delta.content or "")
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

    return full_response

def create_intro(name):
    improve_message = [
        {"role": "system", "content": """You speak the Serbian language and your task is to adapt the sentence I will give you by correcting for Grammar and Gender. 
        Be careful with Serbian names double check if they are male or female names. 
        If the name ends with one of the following letters {a, e, i, o, u} don't change it (e.g. Milica >> Milica)
        If it doesn't end with one of those letters you should add the letter "e" at the end of the name (e.g. Goran >> Gorane), UNLESS if it ends with the letters "ar", in that case replace those last two letters with "re" (e.g. Aleksandar >> Aleksandre)
        DO NOT COMMENT only correct."""},
        {"role": "user", "content": f"Poštovani {name}, izveštaj je u prilogu."}
    ]
    response = client.chat.completions.create(
        model=work_vars["names"]["openai_model"],
        messages=improve_message,
    )
    return response.choices[0].message.content

def recommended(full_response):
    processor = HybridQueryProcessor(namespace="positive", top_k=3)
    return processor.process_query_results(full_response)

def main():
    opcija = st.query_params.get('opcija', "Sve")
    if opcija == "Sve":
        with st.sidebar:
            st.caption("Ver. 30.05.24")
            st.subheader("GAP analiza")
            opcija = st.selectbox("Odaberite upitnik", ("",
                                                        "Opšti",
                                                        "Poslovna zrelost",
                                                        "Digitalna zrelost",
                                                        "Sajber bezbednost",
                                                        "IT infrastruktura",
                                                        "Upotreba AI"))

    if opcija != "":
        result, email, ime = odgovori(opcija)
        file_name = f"{opcija}.docx"
        anketa = format_json_to_text(result).replace('*', '').replace("'", '')

        if result:
            gap_message = [
                {"role": "system", "content": mprompts["gap_ba_expert"]},
                {"role": "user", "content": mprompts["gap_write_report"].format(result=result)}
            ]
            full_response = positive_agent(gap_message)
            predlozi, x, y = recommended(full_response)

            recommend_message = [
                {"role": "system", "content": mprompts["gap_dt_consultant"]},
                {"role": "user", "content": mprompts["gap_service_suggestion"].format(full_response=full_response, predlozi=predlozi)}
            ]
            recommendation_response = positive_agent(recommend_message)

            # gap_analiza = full_response + "\n\n" + recommendation_response
            gap_analiza = recommendation_response
            pdf_file_name = sacuvaj_dokument_upitnik(gap_analiza, file_name, anketa=anketa)
            poruka = create_intro(ime)
            st.info(poruka)
            posalji_mail(email, pdf_file_name, poruka)
            try:
                os.remove(pdf_file_name)
            except:
                pass

if __name__ == "__main__":
    main()
